from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_background(cell, fill_color):
    """
    set_cell_background(cell, "FF0000")
    """
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_full_report():
    doc = Document()

    # Global Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(11)

    # --- TITLE PAGE ---
    doc.add_paragraph('\n' * 2)
    title = doc.add_heading('MICRO-SMART-GRID INTELLIGENT ENERGY MANAGEMENT SYSTEM', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Multi-Scale Performance & Growth Analysis (24h, 1w, 1y)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.color.rgb = RGBColor(46, 116, 181)

    doc.add_paragraph('\n' * 3)
    
    inf_para = doc.add_paragraph()
    inf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = inf_para.add_run('Prepared by: Gemini AI Engineering Team\nDate: April 14, 2026\nProject: Microgrid Optimization V2.0')
    run.font.size = Pt(12)
    run.italic = True

    doc.add_page_break()

    # --- 1. EXECUTIVE OVERVIEW ---
    doc.add_heading('1. Executive Overview', level=1)
    doc.add_paragraph(
        "This comprehensive analysis details the operational efficiency and economic impact of the "
        "Micro-Smart-Grid Management System. By integrating hybrid CNN-LSTM deep learning architectures "
        "with real-time optimization, the system achieves significant grid independence and cost reduction. "
        "This document evaluates performance across three temporal scales: Immediate (24 Hours), "
        "Short-term (1 Week), and Long-term Growth (1 Year)."
    )

    # --- 2. TECHNICAL SPECIFICATIONS ---
    doc.add_heading('2. System Architecture & Technical Specifications', level=1)
    doc.add_paragraph("The HRES (Hybrid Renewable Energy System) is composed of four primary sub-systems:")
    
    tech_specs = [
        ["Component", "Parameter", "Value", "Unit"],
        ["Solar PV", "Rated Capacity", "200", "kW"],
        ["Wind Turbine", "Rated Capacity", "150", "kW"],
        ["Biogas Gen", "Rated Capacity", "100", "kW"],
        ["Battery (BESS)", "Total Capacity", "500", "kWh"]
    ]
    table = doc.add_table(rows=len(tech_specs), cols=len(tech_specs[0]))
    table.style = 'Table Grid'
    for i, row in enumerate(tech_specs):
        for j, val in enumerate(row):
            table.cell(i, j).text = val

    # --- 3. 24-HOUR OPERATIONAL ANALYSIS ---
    doc.add_heading('3. 24-Hour Operational Analysis', level=1)
    doc.add_paragraph(
        "Real-time monitoring of the system over the last 24 hours shows high responsiveness to load fluctuations "
        "and renewable variability. The integration of forecasting allows for pre-emptive battery charging."
    )

    if os.path.exists('plots-24hours/plot_1.png'):
        doc.add_heading('3.1 Generation vs Consumption (24h)', level=2)
        doc.add_picture('plots-24hours/plot_1.png', width=Inches(5.5))
        doc.add_paragraph('Figure 3.1: 24-hour cycle showing peak solar generation and nocturnal wind contribution.', style='Caption')

    if os.path.exists('plots-24hours/plot_2.png'):
        doc.add_heading('3.2 Battery State of Charge (24h)', level=2)
        doc.add_picture('plots-24hours/plot_2.png', width=Inches(5.5))
        doc.add_paragraph('Figure 3.2: SOC management showing efficient energy buffering during daylight hours.', style='Caption')

    doc.add_page_break()

    # --- 4. WEEKLY PERFORMANCE METRICS ---
    doc.add_heading('4. Weekly Performance Metrics', level=1)
    doc.add_paragraph(
        "The system demonstrated 62.03% cost savings over the past week compared to a grid-only baseline."
    )

    # Weekly Table
    weekly_data = [
        ["Date", "Total Ren (kWh)", "Load (kWh)", "Grid (kWh)", "Solar", "Wind", "Biogas"],
        ["2026-04-07", "722.19", "837.33", "0.00", "468.12", "157.59", "96.48"],
        ["2026-04-10", "1776.33", "2886.24", "1092.61", "1167.60", "317.88", "290.85"],
        ["2026-04-13", "1541.64", "2886.24", "1344.91", "1013.78", "185.64", "342.21"],
        ["TOTAL (Week)", "12332.13", "20203.67", "7671.54", "8280.53", "1809.68", "2241.92"]
    ]
    table_w = doc.add_table(rows=len(weekly_data), cols=len(weekly_data[0]))
    table_w.style = 'Table Grid'
    for i, row in enumerate(weekly_data):
        for j, val in enumerate(row):
            table_w.cell(i, j).text = val

    if os.path.exists('plots-1week/plot_4.png'):
        doc.add_heading('4.1 Total Gen vs Load (Weekly Trend)', level=2)
        doc.add_picture('plots-1week/plot_4.png', width=Inches(5.5))
        doc.add_paragraph('Figure 4.1: Weekly aggregate of renewable vs load demand.', style='Caption')

    # --- 5. LONG-TERM GROWTH (YEARLY ANALYSIS) ---
    doc.add_page_break()
    doc.add_heading('5. Long-Term Growth (1-Year Analysis)', level=1)
    doc.add_paragraph(
        "Over the past 12 months, the system has scaled effectively to meet seasonal demand changes. "
        "Total renewable production reached 708,476 kWh, offsetting approximately 67.5% of the total yearly load."
    )

    # Monthly Growth Table
    monthly_data = [
        ["Month", "Total Ren (kWh)", "Load (kWh)", "Solar (kWh)", "Wind (kWh)", "Biogas (kWh)"],
        ["2025-06", "57159", "91176", "41525", "6075", "9558"],
        ["2025-09", "73973", "68675", "49537", "14972", "9462"],
        ["2025-12", "57833", "84825", "38295", "10212", "9325"],
        ["2026-03", "48867", "106717", "33136", "5663", "10067"],
        ["TOTAL", "708476", "1050985", "485274", "105023", "118178"]
    ]
    table_m = doc.add_table(rows=len(monthly_data), cols=len(monthly_data[0]))
    table_m.style = 'Table Grid'
    for i, row in enumerate(monthly_data):
        for j, val in enumerate(row):
            table_m.cell(i, j).text = val

    if os.path.exists('plots-1year/plot_1.png'):
        doc.add_heading('5.1 Seasonal Generation Profile', level=2)
        doc.add_picture('plots-1year/plot_1.png', width=Inches(5.5))
        doc.add_paragraph('Figure 5.1: Annual generation profile highlighting seasonal solar and wind peaks.', style='Caption')

    if os.path.exists('plots-1year/plot_10.png'):
        doc.add_heading('5.2 Yearly Solar Growth & Performance', level=2)
        doc.add_picture('plots-1year/plot_10.png', width=Inches(5.5))
        doc.add_paragraph('Figure 5.2: Year-long solar forecasting accuracy and production trends.', style='Caption')

    # --- 6. MODEL PERFORMANCE ---
    doc.add_page_break()
    doc.add_heading('6. Predictive Model Performance', level=1)
    doc.add_paragraph(
        "The proposed CNN-LSTM model outperformed standard architectures. The hybrid approach "
        "captures both spatial features (via CNN) and temporal dependencies (via LSTM)."
    )
    
    perf_metrics = [
        ["Method", "MAE", "RMSE", "R2 Score"],
        ["CNN-LSTM (Proposed)", "0.0928", "0.1257", "0.2494"],
        ["LSTM", "0.0973", "0.1284", "0.2230"],
        ["ARIMA (Baseline)", "0.1820", "0.2100", "0.1100"]
    ]
    table_p = doc.add_table(rows=len(perf_metrics), cols=len(perf_metrics[0]))
    table_p.style = 'Table Grid'
    for i, row in enumerate(perf_metrics):
        for j, val in enumerate(row):
            table_p.cell(i, j).text = val

    # --- 7. COST & ECONOMIC IMPACT SUMMARY ---
    doc.add_heading('7. Final Economic Impact Summary', level=1)
    
    grid_cost = 161629.36
    optimized_cost = 61372.32
    savings = grid_cost - optimized_cost

    doc.add_paragraph(f"Net Weekly Savings: \u20b9 {savings:,.2f}")
    doc.add_paragraph(f"Yearly Projected Savings: \u20b9 {savings * 52:,.2f}")
    
    if os.path.exists('plots-1week/plot_5.png'):
        doc.add_picture('plots-1week/plot_5.png', width=Inches(3.5))
        doc.add_paragraph('Figure 7.1: Comparison of operational costs.', style='Caption')

    # Conclusion
    doc.add_heading('8. Conclusion', level=1)
    doc.add_paragraph(
        "The Micro-Smart-Grid system demonstrates consistent high performance across all evaluated timeframes. "
        "The 1-year analysis confirms the system's resilience, while the 24-hour and weekly data validate its "
        "operational precision. Future expansions will consider incorporating green hydrogen storage to "
        "further reduce grid reliance during low renewable production months."
    )

    doc.save('Microgrid_Full_Analysis_Report.docx')
    print("Full report created: Microgrid_Full_Analysis_Report.docx")

if __name__ == "__main__":
    create_full_report()
