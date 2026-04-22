from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_report():
    doc = Document()

    # Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Title Page
    title = doc.add_heading('Micro-Smart-Grid Energy Management System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Comprehensive Performance Analysis Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(16)

    doc.add_paragraph('\n' * 5)
    
    date_para = doc.add_paragraph('Date: April 14, 2026')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    # Introduction
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "This report presents a detailed analysis of the Micro-Smart-Grid Energy Management System. "
        "The system leverages deep learning models (LSTM and CNN-LSTM) for precise energy forecasting "
        "of renewable sources, including solar, wind, and biomass. An intelligent optimization module "
        "is employed to manage energy storage and minimize grid dependency, resulting in significant cost savings."
    )

    doc.add_heading('2. System Features', level=2)
    features = [
        "Multi-Source Forecasting: Deep learning models to predict energy generation.",
        "Microgrid Optimization: Intelligent management of battery storage and grid interaction.",
        "Model Comparison: Continuous evaluation of neural network architectures.",
        "Visualization: Real-time graphical summaries of forecasting and performance."
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')

    # Weekly Performance
    doc.add_heading('3. Weekly Performance Analysis', level=1)
    doc.add_paragraph(
        "The following analysis summarizes the performance of the microgrid over the past week (April 7 - April 14, 2026). "
        "The system demonstrates robust energy balancing between renewable generation and varying load demands."
    )

    # Table 1: Daily Performance
    doc.add_heading('3.1 Daily Generation and Load Summary', level=2)
    table_data = [
        ["Date", "Total Ren (kWh)", "Load (kWh)", "Grid (kWh)", "Solar (kWh)", "Wind (kWh)", "Biogas (kWh)"],
        ["2026-04-07", "722.19", "837.33", "0.00", "468.12", "157.59", "96.48"],
        ["2026-04-08", "1790.49", "2886.24", "1010.89", "1124.58", "351.60", "314.31"],
        ["2026-04-09", "1929.40", "2886.24", "974.14", "1312.13", "281.85", "335.42"],
        ["2026-04-10", "1776.33", "2886.24", "1092.61", "1167.60", "317.88", "290.85"],
        ["2026-04-11", "1657.15", "2886.24", "1229.09", "1116.53", "224.86", "315.76"],
        ["2026-04-12", "1805.40", "2886.24", "1081.96", "1334.27", "141.03", "330.10"],
        ["2026-04-13", "1541.64", "2886.24", "1344.91", "1013.78", "185.64", "342.21"],
        ["2026-04-14", "1109.53", "2048.91", "937.94", "743.52", "149.23", "216.78"],
        ["TOTAL", "12332.13", "20203.67", "7671.54", "8280.53", "1809.68", "2241.92"]
    ]

    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = 'Table Grid'
    for i, row in enumerate(table_data):
        for j, val in enumerate(row):
            table.cell(i, j).text = val

    doc.add_paragraph('\n')
    
    # Plot 1: Generation vs Consumption
    if os.path.exists('plots-1week/plot_1.png'):
        doc.add_heading('3.2 Energy Generation vs. Consumption Trends', level=2)
        doc.add_picture('plots-1week/plot_1.png', width=Inches(6))
        doc.add_paragraph('Figure 1: Comparison of total renewable generation versus load demand over 7 days.')

    doc.add_page_break()

    # Energy Storage
    doc.add_heading('4. Energy Storage and Battery Optimization', level=1)
    doc.add_paragraph(
        "The battery storage system plays a critical role in balancing supply and demand. "
        "The optimization module manages the State of Charge (SOC) to maximize self-consumption of renewable energy."
    )

    if os.path.exists('plots-1week/plot_2.png'):
        doc.add_heading('4.1 Battery State of Charge (SOC)', level=2)
        doc.add_picture('plots-1week/plot_2.png', width=Inches(6))
        doc.add_paragraph('Figure 2: Battery SOC dynamics during the weekly cycle.')

    # Table 4: Sensitivity
    doc.add_heading('4.2 Sensitivity Analysis: Battery Capacity Impact', level=2)
    sens_data = [
        ["Battery Capacity (kWh)", "Grid Purchase (kWh)", "Total Cost (\u20b9)", "Cost Reduction (%)"],
        ["250", "7901.96", "63215.69", "60.89%"],
        ["500", "7671.54", "61372.33", "62.03%"],
        ["750", "7671.54", "61372.33", "62.03%"],
        ["1000", "7671.54", "61372.33", "62.03%"],
        ["1500", "7671.54", "61372.33", "62.03%"]
    ]
    table_sens = doc.add_table(rows=len(sens_data), cols=len(sens_data[0]))
    table_sens.style = 'Table Grid'
    for i, row in enumerate(sens_data):
        for j, val in enumerate(row):
            table_sens.cell(i, j).text = val

    doc.add_page_break()

    # Forecasting
    doc.add_heading('5. Renewable Energy Forecasting Accuracy', level=1)
    doc.add_paragraph(
        "Accurate forecasting is essential for optimal dispatch. The LSTM/CNN-LSTM models provide high-precision "
        "short-term predictions for solar, wind, and biogas generation."
    )

    if os.path.exists('plots-1week/plot_7_1.png'):
        doc.add_heading('5.1 Solar Generation Forecasting', level=2)
        doc.add_picture('plots-1week/plot_7_1.png', width=Inches(6))
        doc.add_paragraph('Figure 3: Predicted vs. Actual Solar Generation.')

    if os.path.exists('plots-1week/plot_9.png'):
        doc.add_heading('5.2 Wind Generation Trends', level=2)
        doc.add_picture('plots-1week/plot_9.png', width=Inches(6))
        doc.add_paragraph('Figure 4: Predicted vs. Actual Wind Generation.')

    doc.add_page_break()

    # Cost Analysis
    doc.add_heading('6. Economic Impact and Cost Analysis', level=1)
    doc.add_paragraph(
        "By prioritizing renewable sources and optimizing battery usage, the system significantly reduces grid costs."
    )

    cost_data = [
        ["Scenario", "Total Load (kWh)", "Grid Usage (kWh)", "Unit Cost (\u20b9)", "Total Cost (\u20b9)"],
        ["Grid-only System", "20,203.67", "20,203.67", "8.00", "161,629.36"],
        ["Hybrid (Optimized)", "20,203.67", "7,671.54", "8.00", "61,372.32"]
    ]
    table_cost = doc.add_table(rows=len(cost_data), cols=len(cost_data[0]))
    table_cost.style = 'Table Grid'
    for i, row in enumerate(cost_data):
        for j, val in enumerate(row):
            table_cost.cell(i, j).text = val

    doc.add_paragraph('\n')
    doc.add_paragraph('NET WEEKLY SAVINGS: \u20b9 100,257.04', style='Heading 2')
    doc.add_paragraph('PERCENTAGE SAVINGS: 62.03%', style='Heading 2')

    if os.path.exists('plots-1week/plot_5.png'):
        doc.add_picture('plots-1week/plot_5.png', width=Inches(4))
        doc.add_paragraph('Figure 5: Cost comparison: Grid-only vs. Optimized Hybrid System.')

    # Conclusion
    doc.add_heading('7. Conclusion', level=1)
    doc.add_paragraph(
        "The Micro-Smart-Grid Energy Management System has proven highly effective during the analysis period. "
        "With a cost reduction of over 62%, the system demonstrates the economic viability of integrated renewable energy "
        "management. Future improvements will focus on refining forecasting models for even higher accuracy during extreme weather events."
    )

    doc.save('Microgrid_Performance_Report.docx')
    print("Report created: Microgrid_Performance_Report.docx")

if __name__ == "__main__":
    create_report()
